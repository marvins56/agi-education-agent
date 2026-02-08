"""DOCX document parser using python-docx.

Extracts paragraphs with heading hierarchy, tables as markdown,
and header/footer text from DOCX files.
"""

import logging

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

# Mapping of python-docx heading style names to markdown heading levels
_HEADING_LEVEL_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
    "Subtitle": 2,
}


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to a markdown-formatted table string."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Replace pipes in cell text to avoid breaking markdown tables
            cell_text = cell.text.strip().replace("|", "\\|")
            cells.append(cell_text)
        rows.append(cells)

    if not rows:
        return ""

    lines = []
    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    # Separator row
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row if it has fewer cells than header
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

    return "\n".join(lines)


def _extract_header_footer_text(doc) -> tuple[list[str], list[str]]:
    """Extract header and footer text from all sections of the document."""
    header_texts = []
    footer_texts = []

    for section in doc.sections:
        try:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        header_texts.append(text)
        except Exception as exc:
            logger.debug("Could not read header in section: %s", exc)

        try:
            if section.footer and section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        footer_texts.append(text)
        except Exception as exc:
            logger.debug("Could not read footer in section: %s", exc)

    # Deduplicate while preserving order
    seen_h: set[str] = set()
    unique_headers = []
    for t in header_texts:
        if t not in seen_h:
            seen_h.add(t)
            unique_headers.append(t)

    seen_f: set[str] = set()
    unique_footers = []
    for t in footer_texts:
        if t not in seen_f:
            seen_f.add(t)
            unique_footers.append(t)

    return unique_headers, unique_footers


class DocxParser:
    """Extract text and metadata from DOCX files.

    Supports paragraph text with heading hierarchy, tables formatted
    as markdown, and header/footer extraction.
    """

    def parse(self, file_path: str) -> str:
        """Extract all text from a DOCX file.

        Returns plain text with headings prefixed by # markers,
        tables formatted as markdown, and header/footer text appended.
        """
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            logger.error("File is not a valid DOCX package: %s", file_path)
            return ""
        except Exception as exc:
            logger.error("Failed to open DOCX file %s: %s", file_path, exc)
            return ""

        parts = self._extract_body(doc)

        # Append headers/footers
        headers, footers = _extract_header_footer_text(doc)
        if headers:
            parts.append("\n[Headers]")
            parts.extend(headers)
        if footers:
            parts.append("\n[Footers]")
            parts.extend(footers)

        return "\n\n".join(parts)

    def parse_with_metadata(self, file_path: str) -> dict:
        """Extract text and metadata from a DOCX file."""
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            logger.error("File is not a valid DOCX package: %s", file_path)
            return {"text": "", "pages": 0, "metadata": {}}
        except Exception as exc:
            logger.error("Failed to open DOCX file %s: %s", file_path, exc)
            return {"text": "", "pages": 0, "metadata": {}}

        parts = self._extract_body(doc)
        heading_count = self._count_headings(doc)
        table_count = len(doc.tables)

        # Append headers/footers
        headers, footers = _extract_header_footer_text(doc)
        if headers:
            parts.append("\n[Headers]")
            parts.extend(headers)
        if footers:
            parts.append("\n[Footers]")
            parts.extend(footers)

        text = "\n\n".join(parts)

        try:
            core = doc.core_properties
            meta = {
                "title": core.title or "",
                "author": core.author or "",
                "subject": core.subject or "",
                "table_count": table_count,
                "heading_count": heading_count,
            }
        except Exception as exc:
            logger.warning("Could not read core properties: %s", exc)
            meta = {
                "table_count": table_count,
                "heading_count": heading_count,
            }

        return {
            "text": text,
            "pages": len(doc.sections),
            "metadata": meta,
        }

    def _extract_body(self, doc) -> list[str]:
        """Extract paragraphs and tables in document order.

        Paragraphs with heading styles are prefixed with markdown heading markers.
        Tables are converted to markdown table format.
        """
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            try:
                text = paragraph.text.strip()
                if not text:
                    continue

                style_name = paragraph.style.name if paragraph.style else ""
                heading_level = _HEADING_LEVEL_MAP.get(style_name, 0)

                if heading_level > 0:
                    parts.append(f"{'#' * heading_level} {text}")
                else:
                    parts.append(text)
            except Exception as exc:
                logger.debug("Error processing paragraph: %s", exc)
                continue

        # Append tables separately (interleaving with exact paragraph position
        # would require XML-level iteration which is fragile; appending after
        # paragraphs is the pragmatic approach used by most DOCX extractors).
        for idx, table in enumerate(doc.tables):
            try:
                md_table = _table_to_markdown(table)
                if md_table:
                    parts.append(md_table)
            except Exception as exc:
                logger.debug("Error processing table %d: %s", idx, exc)
                continue

        return parts

    def _count_headings(self, doc) -> int:
        """Count paragraphs with heading styles."""
        count = 0
        for paragraph in doc.paragraphs:
            try:
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name in _HEADING_LEVEL_MAP:
                    count += 1
            except Exception:
                continue
        return count
